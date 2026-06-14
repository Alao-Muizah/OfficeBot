import os
import csv 
import openpyxl
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer 
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
import pdfplumber




# =========== Getting the absolute file path to use =========== #

def _resolve_path(path: str, workdir:str) -> str:
    if os.path.isabs(path) :
        return path
    return os.path.join(workdir, path)

# =========== Converting word doc to pdf file =========== #
def convert_word_to_pdf(input_path: str, output_path: str, workdir:str) -> dict:

    full_input = _resolve_path(input_path, workdir)
    full_output = _resolve_path(output_path, workdir)

    if not os.path.exists(full_input):
        return {"status": "error", "message": f"File not found: {full_input}"}
    
    doc = Document(full_input)

    Paragraph_text = [p.text for p in doc.paragraphs if p.text.strip()]
    if not Paragraph_text:
        return {"status": "error", "message": "Word document is empty, nothing to convert"}
    
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    heading_style = styles['Heading1']

    pdf = SimpleDocTemplate(
        full_output,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm 
    )

    flowables = []

    for i, text in enumerate(Paragraph_text):

        style = heading_style if i == 0 else normal_style
        flowables.append(Paragraph(text, style))
        flowables.append(Spacer(1, 0.3 * cm))

    pdf.build(flowables)

    return {
        "status": "ok",
        "message": f"Converted '{input_path}' to PDF",
        "output": full_output
    }


# =========== Converting PDF to Word =========== #

def convert_pdf_to_word(input_path: str, output_path: str, workdir: str) -> dict:

    full_input = _resolve_path(input_path, workdir)
    full_output = _resolve_path(output_path, workdir)

    if not os.path.exists(full_input):
        return {
            "status": "error",
            "message": f"File not found: {full_input}"
        }
    
    with pdfplumber.open(full_input) as pdf:

        doc = Document()

        for page_num, page in enumerate(pdf.pages, start=1):

            text = page.extract_text()

            if text:
                doc.add_heading(f"Page {page_num}", level=2)

                for line in text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
            else:
                doc.add_paragraph(f"[Page {page_num}: no extractable text]")

    doc.save(full_output)

    return {
        "staus": "ok",
        "message": f"Converted '{input_path}' to Word document",
        "output": full_output
    }


# =========== Converting excel to csv =========== #

def convert_excel_to_csv(input_path: str, output_path: str, sheet_name: str, workdir: str) -> dict:

    full_input = _resolve_path(input_path, workdir)
    full_output = _resolve_path(output_path, workdir)

    if not os.path.exists(full_input):
        return {
            "status": "error",
            "message": f"File not found {full_input}"
        }
    
    wb = openpyxl.load_workbook(full_input)

    if sheet_name not in wb.sheetnames:
        return {
            "status": "error",
            "message": f"File not found {full_input}"
        }
    
    ws = wb[sheet_name]

    all_rows = list(ws.iter_rows(values_only=True))

    if not all_rows:
        return {
            "status":"error",
            "message": f"Sheet {sheet_name} is empty, Nothing to convert"
        }
    
    with open(full_output, "w", newline="", encoding="utf-8-sig") as f:

        writer = csv.writer(f)

        writer.writerows(all_rows)

    row_count = len(all_rows) -1 

    return {
        "status": "ok",
        "message": f"Converted sheet '{sheet_name}' t CSV ({row_count}) data rows",
        "output": full_output
    }


# =========== Converting CSV to excel =========== #

def convert_csv_to_excel(input_path: str, output_path: str, sheet_name: str, workdir:str) -> dict:

    full_input = _resolve_path(input_path, workdir)
    full_output = _resolve_path(output_path, workdir)

    if not os.path.exists(full_input):
        return {
            "status": "error",
            "message": F"File not found {sheet_name}"
        }
    
    with open(full_input, "r", newline="", encoding="utf-8-sig") as f:

        reader = csv.reader(f)
        all_rows = list(reader)

        if not all_rows:
            return {
                "status": "error",
                "message": f"CSV file '{full_input}' is empty"
            }
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        for row in all_rows:

            converted_row = []

            for value in row:

                try:
                    converted_row.append(int(value))
                except ValueError:
                    
                    try :
                        converted_row.append(float(value))
                    except ValueError:

                        converted_row.append(value)

            ws.append(converted_row)
        
        wb.save(full_output)

        data_rows = len(all_rows) - 1
        return {
            "status": "ok",
            "message": f"Converted CSV to Excel sheet'{full_input}' ({data_rows} data rows)",
            "output": full_output
        }
    
