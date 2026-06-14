from docx import Document
from docx.shared import Pt
import os 
from typing import Literal


# =========== Getting the absolute file path to use =========== #

def _resolve_path(path: str, workdir: str) -> str:
    if os.path.isabs(path):
        return path
    return os.path.join(workdir, path)

# ====================== Creating the word doc ====================== #

def create_word_doc(path: str, title: str, body: str, workdir: str) -> dict:
    full_path = _resolve_path(path, workdir)

    parent_dir = os.path.dirname(full_path)
    if parent_dir:
        os.makedirs(parent_dir, exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)
    doc.save(full_path)

    return {"status": "ok", 
            "message": f"Created '{path}'", 
            "path": full_path}


# ====================== Reading the word doc ====================== #

def read_word_doc(path:str, workdir:str)-> dict:

    full_path = _resolve_path(path, workdir)
    if not os.path.exists(full_path):
        return {"status": "error", "message": f"file not found {full_path}"}
    
    doc = Document(full_path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
 
    content = "\n".join(lines)
 
    return {"status": "ok", "content": content, "paragraph_count": len(lines)}


# ====================== Append new paragraph ====================== #

def append_paragraph(path: str, text: str, workdir: str) -> dict:
    full_path = _resolve_path(path, workdir)
 
    if not os.path.exists(full_path):
        return {"status": "error", "message": f"File not found: {full_path}"}
 
    doc = Document(full_path)
    doc.add_paragraph(text)
    doc.save(full_path)
 
    return {"status": "ok",
             "message": f"Appended paragraph to '{path}'"}


# ====================== Add Heading ====================== #

def add_heading(path: str, text: str, level: int, workdir: str) -> dict:
   
    full_path = _resolve_path(path, workdir)
 
    if not os.path.exists(full_path):
        return {"status": "error", "message": f"File not found: {full_path}"}
 
    doc = Document(full_path)
    doc.add_heading(text, level=level)
    doc.save(full_path)
 
    return {"status": "ok", "message": f"Added heading '{text}' (level {level})"}


# ====================== Replacing text ====================== #

def replace_text(path:str, workdir:str, old_text:str, new_text:str)-> dict:

    full_path = _resolve_path(path, workdir) 

    if not os.path.exists(full_path):
        return {"status": "error",
                 "message": f"File not found: {full_path}"}

    doc = Document(full_path)
    count = 0

    for para in doc.paragraphs:

        if old_text in para.text:

            new_full_text = para.text.replace(old_text, new_text)
            count += 1
            
            runs_to_remove = list(para.runs)
            for run in runs_to_remove:
                para._p.remove(run._r)

            para.add_run(new_full_text)

    doc.save(full_path)

    return {"status": "ok",
             "message": f"Replaced '{old_text}' with '{new_text}' in {count} paragraph(s)"}


# ====================== Set/Change font ====================== #

def set_font(path: str, font_name: str, font_size: int, workdir: str) -> dict:
    full_path = _resolve_path(path, workdir)

    if not os.path.exists(full_path):
        return {"status": "error", "message": f"File not found: {full_path}"}

    doc = Document(full_path)

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    doc.save(full_path)

    return {"status": "ok", "message": f"Set font to {font_name} {font_size}pt across entire document"}

# ====================== Insert table ====================== #

def insert_table(path: str, headers: list, rows: list, workdir: str) -> dict:
    full_path = _resolve_path(path, workdir)

    if not os.path.exists(full_path):
        return {"status": "error", "message": f"File not found: {full_path}"}

    doc = Document(full_path)

    col_count = len(headers)
    row_count = len(rows) + 1

    table = doc.add_table(rows=row_count, cols=col_count)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = str(value)

    doc.save(full_path)

    return {
        "status": "ok",
        "message": f"Inserted table with {len(headers)} columns and {len(rows)} data rows"
    }

# ====================== Text Formatting ====================== #


def text_format(path:str, target_text:str, formatting:Literal['bold', 'italic', 'underlined'] , workdir:str) -> dict:

    if formatting not in ('bold', 'italic', 'underlined'):
        raise ValueError("text format must be bold, text or underlined") 
    
    full_path = _resolve_path(path, workdir)
    if not os.path.exists(full_path):
        return {"status": "error", "message": f"File not found: {full_path}"}
    
    doc = Document(full_path)
    count = 0

    for para in doc.paragraphs:

        if target_text in para.text:
            full_text = para.text

            before, match, after = full_text.partition(target_text)

            runs_to_remove = list(para.runs)
            for run in runs_to_remove:
                para._p.remove(run._r)

            if before:
                para.add_run(before)   

            text_run = para.add_run(match) 
            if formatting == 'bold': 
                text_run.bold = True
            elif formatting == 'italic': 
                text_run.italic = True
            elif formatting == 'underlined': 
                text_run.underline = True

            if after:
                para.add_run(after)          
 
            count += 1
        else:
            return {"status": "failed",
                "message": f"'{target_text}' not found in {full_path}"} 
    
    doc.save(full_path)

    return {"status": "ok",
             "message": f"Bolded '{target_text}' in {count} paragraph(s)"}


